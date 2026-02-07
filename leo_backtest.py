"""
Backtest three moving average crossover strategies (EMA 9/20, EMA 20/50, SMA 20/50)
on a watchlist of stocks using yfinance data. Generates equity curve plots, a summary CSV,
and a Word document report with results.

Author: Leonardo Carcache
Last Update: January 13, 2026
"""


#!/usr/bin/env python3
import pandas as pd
import numpy as np
import yfinance as yf
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document
import os
import time
import traceback

# Set matplotlib style
plt.style.use('seaborn-v0_8')

# Paths (relative)
watchlist_csv = 'watchlist_leo_project7.csv'   # CSV must be in same folder as this script
output_dir = './outputs'
os.makedirs(output_dir, exist_ok=True)

# Change to True to include the Summary Table in the Word report that is generated in the outputs folder.
include_table = False

# Load watchlist
try:
    watchlist_df = pd.read_csv(watchlist_csv)
except FileNotFoundError:
    print(f"ERROR: Watchlist CSV not found at {watchlist_csv}. Place the CSV in the same folder as this script.")
    raise

if 'Symbol' not in watchlist_df.columns:
    print("ERROR: CSV must contain a column named 'Symbol'.")
    raise SystemExit(1)

symbols = watchlist_df['Symbol'].dropna().unique().tolist()

# Backtest parameters
# The backtest runs from Jan 1, 2020 to today.
initial_capital = 10000.0
end_date = datetime.today()
start_date = datetime(2020, 1, 1)

print("=== Debug Header ===")
print(f"Working directory: {os.getcwd()}")
print(f"Backtest period: {start_date.date()} to {end_date.date()}")
print(f"Symbols to process: {len(symbols)}")
print("====================\n")

# Metrics storage
results = []
failed_symbols = []
global_trade_log = []

def compute_indicators(df):
    # EMAs
    df['EMA9'] = df['Adj Close'].ewm(span=9, adjust=False).mean()
    df['EMA20'] = df['Adj Close'].ewm(span=20, adjust=False).mean()
    df['EMA30'] = df['Adj Close'].ewm(span=30, adjust=False).mean()
    df['EMA50'] = df['Adj Close'].ewm(span=50, adjust=False).mean()
    df['EMA60'] = df['Adj Close'].ewm(span=60, adjust=False).mean()
    df['EMA200'] = df['Adj Close'].ewm(span=200, adjust=False).mean()
    
    # SMAs
    df['SMA9'] = df['Adj Close'].rolling(window=9).mean()
    df['SMA20'] = df['Adj Close'].rolling(window=20).mean()
    df['SMA30'] = df['Adj Close'].rolling(window=30).mean()
    df['SMA50'] = df['Adj Close'].rolling(window=50).mean()
    df['SMA60'] = df['Adj Close'].rolling(window=60).mean()
    df['SMA200'] = df['Adj Close'].rolling(window=200).mean()
    return df

def generate_signals(df, short_col, long_col):
    cond = (df[short_col] > df[long_col]).astype(int)
    sig = cond.diff().fillna(0)
    return sig

def backtest_strategy(df, short_col, long_col, symbol, strategy_name):
    df = df.copy()
    df['Signal'] = generate_signals(df, short_col, long_col)
    position = 0
    cash = initial_capital
    shares = 0.0
    equity_curve = []
    trades = []

    for i in range(len(df)):
        price = df['Adj Close'].iloc[i]
        signal = df['Signal'].iloc[i]
        date = df.index[i]

        if signal == 1 and position == 0 and price > 0:
            shares = cash / price
            cash = 0.0
            position = 1
            trades.append({'entry_date': date, 'entry_price': price})
            
            # Log BUY
            global_trade_log.append({
                'Symbol': symbol,
                'Strategy': strategy_name,
                'Date': date.strftime('%Y-%m-%d'),
                'Action': 'BUY',
                'Price': price,
                'Shares': shares
            })

        elif signal == -1 and position == 1:
            cash = shares * price
            position = 0
            trades[-1].update({'exit_date': date, 'exit_price': price})
            
            # Log SELL
            global_trade_log.append({
                'Symbol': symbol,
                'Strategy': strategy_name,
                'Date': date.strftime('%Y-%m-%d'),
                'Action': 'SELL',
                'Price': price,
                'Shares': shares
            })
            shares = 0.0

        equity = cash + shares * price
        equity_curve.append(equity)

    # Force close at end
    if position == 1:
        last_price = df['Adj Close'].iloc[-1]
        cash = shares * last_price
        trades[-1].update({'exit_date': df.index[-1], 'exit_price': last_price})
        
        # Log final SELL (forced)
        global_trade_log.append({
            'Symbol': symbol,
            'Strategy': strategy_name,
            'Date': df.index[-1].strftime('%Y-%m-%d'),
            'Action': 'SELL (Forced)',
            'Price': last_price,
            'Shares': shares
        })
        shares = 0.0

    final_value = cash
    returns = (final_value - initial_capital) / initial_capital
    days = (df.index[-1] - df.index[0]).days
    years = days / 365.25 if days > 0 else 1.0
    cagr = (final_value / initial_capital) ** (1 / years) - 1 if years > 0 else 0.0

    equity_series = pd.Series(equity_curve, index=df.index)
    peak = equity_series.cummax()
    drawdown = (equity_series - peak) / peak
    max_dd = drawdown.min() if not drawdown.empty else 0.0
    
    num_trades = len(trades)
    wins = [1 for t in trades if t.get('exit_price', 0) > t.get('entry_price', 0)]
    win_rate = len(wins) / num_trades if num_trades > 0 else 0.0
    return final_value, returns, cagr, max_dd, num_trades, win_rate, equity_series

# Define strategies
strategies = [
    ('EMA', 9, 20), ('EMA', 9, 50), ('EMA', 20, 50), ('EMA', 30, 60), ('EMA', 50, 200),
    ('SMA', 9, 20), ('SMA', 9, 50), ('SMA', 20, 50), ('SMA', 30, 60), ('SMA', 50, 200)
]

for idx, symbol in enumerate(symbols, start=1):
    try:
        print(f"[{idx}/{len(symbols)}] Downloading {symbol} ...")
        df = None
        for attempt in range(2):
            try:
                df = yf.download(symbol, start=start_date, end=end_date, progress=False)
                if df is not None and not df.empty:
                    break
            except Exception:
                time.sleep(1.0)
        time.sleep(0.4)

        if df is None:
            print(f"  -> yfinance returned None for {symbol}, skipping.")
            failed_symbols.append(symbol)
            continue

        # Clean/Flatten columns as before
        if hasattr(df.columns, 'nlevels') and df.columns.nlevels > 1:
            df.columns = df.columns.get_level_values(0)
        else:
            df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]

        if df.empty:
            print(f"  -> No data for {symbol}, skipping.")
            failed_symbols.append(symbol)
            continue

        if 'Adj Close' not in df.columns:
            if 'Close' in df.columns:
                df['Adj Close'] = df['Close']
                print(f"  -> 'Adj Close' missing for {symbol}, using 'Close' as fallback.")
            else:
                print(f"  -> No price column for {symbol}, skipping.")
                failed_symbols.append(symbol)
                continue

        if not isinstance(df.index, pd.DatetimeIndex):
            df.index = pd.to_datetime(df.index)

        df = compute_indicators(df)

        # Ensure we have data for the longest lookback (200)
        # We need all indicators to be present to run all strategies fairly on the same dataset?
        # Or at least for the specific strategy working.
        # To simplify, we dropna based on the longest required period (200).
        required = ['EMA200', 'SMA200'] # These are the longest
        df = df.dropna(subset=required)

        if df.empty:
            print(f"  -> Not enough rows after indicator warm-up (need >200 days) for {symbol}, skipping.")
            failed_symbols.append(symbol)
            continue

        print(f"  -> Indicator sample for {symbol} (Last row):\n{df[['Adj Close','EMA50','SMA200']].tail(1)}")

        # Run all strategies
        symbol_results = [symbol]
        equity_curves = {}

        for strat_type, short, long in strategies:
            strat_name = f"{strat_type}{short}_{long}"
            short_col = f"{strat_type}{short}"
            long_col = f"{strat_type}{long}"
            
            final, ret, cagr, dd, trades, win, eq = backtest_strategy(df, short_col, long_col, symbol, strat_name)
            symbol_results.append(final) # Key requirement: Final Value
            # We could append other metrics, but the user specifically asked for a table with "Symbol, EMA(9,20) final value..."
            # Let's verify if we need detailed metrics for all. The prompt says "Generate a CSV table... columns: Symbol... FinalValue".
            # I will stick to just FinalValue for readability as requested, or maybe minimal extra.
            # User output requirements: "CSV summary table ... columns: Symbol, EMA(9,20) final value, etc."
            
            equity_curves[strat_name] = eq
            
        results.append(symbol_results)

        # Plotting - Keep minimal (Original 3 only)
        # Original 3: EMA 9/20, EMA 20/50, SMA 20/50
        plt.figure(figsize=(10, 6))
        if 'EMA9_20' in equity_curves: plt.plot(equity_curves['EMA9_20'], label='EMA 9/20')
        if 'EMA20_50' in equity_curves: plt.plot(equity_curves['EMA20_50'], label='EMA 20/50')
        if 'SMA20_50' in equity_curves: plt.plot(equity_curves['SMA20_50'], label='SMA 20/50')
        
        plt.title(f'Equity Curve - {symbol} (Selected Strategies)')
        plt.xlabel('Date')
        plt.ylabel('Portfolio Value ($)')
        plt.legend()
        plt.tight_layout()
        plot_path = os.path.join(output_dir, f'{symbol}_equity.png')
        plt.savefig(plot_path)
        plt.close()
        print(f"  -> Done. Saved plot to {plot_path}\n")

    except Exception as e:
        print(f"  -> Failed {symbol}: {e}")
        traceback.print_exc()
        failed_symbols.append(symbol)

# 1. Summary CSV
# Columns: Symbol + 10 FinalValues
strat_names = [f"{t}{s}_{l}" for t, s, l in strategies]
columns = ['Symbol'] + [f"FinalValue_{name}" for name in strat_names]

summary_df = pd.DataFrame(results, columns=columns)
csv_path = os.path.join(output_dir, 'backtest_summary.csv')
summary_df.to_csv(csv_path, index=False)
print(f"Summary CSV saved to {csv_path}")

# 2. Trade Log CSV
trade_log_df = pd.DataFrame(global_trade_log)
if not trade_log_df.empty:
    # Reorder columns
    cols = ['Symbol', 'Strategy', 'Date', 'Action', 'Price', 'Shares']
    trade_log_df = trade_log_df[cols]
    log_path = os.path.join(output_dir, 'trade_log.csv')
    trade_log_df.to_csv(log_path, index=False)
    print(f"Trade Log CSV saved to {log_path}")
else:
    print("No trades generated.")

# 3. Word Report
doc = Document()
doc.add_heading('Backtest Report: Project 8 Extended', 0)
doc.add_paragraph(f'Backtest Period: {start_date.date()} to {end_date.date()}')
doc.add_paragraph('Initial Capital: $10,000 per strategy per symbol')
doc.add_paragraph('Strategies Tested (10 total):', style='List Bullet')
for s in strat_names:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Summary Table (First 10 rows)', level=1)
if not summary_df.empty:
    # Create valid table
    table = doc.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.replace('FinalValue_', '') # Shorten header
        
    for _, row in summary_df.head(10).iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float):
                row_cells[i].text = f'{val:.2f}'
            else:
                row_cells[i].text = str(val)

doc.add_heading('Conclusion', level=1)
doc.add_paragraph("Extended backtest complete. See 'backtest_summary.csv' for full results and 'trade_log.csv' for trade execution details.")

if failed_symbols:
    doc.add_heading('Symbols with Insufficient Data', level=1)
    doc.add_paragraph(', '.join(failed_symbols))

docx_path = os.path.join(output_dir, 'backtest_report.docx')
doc.save(docx_path)
print(f"Word report saved to {docx_path}")

print("Backtest finished.")
